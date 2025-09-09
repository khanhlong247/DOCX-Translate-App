# UI
import sys
from io import BytesIO
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QSplitter, QPushButton,
    QComboBox, QLabel, QHBoxLayout, QFileDialog, QMessageBox
)
from PyQt5.QtWebEngineWidgets import QWebEngineView
from PyQt5.QtCore import QUrl
from docx import Document
from translator_columns import TranslatorColumns
from utils import replace_text_in_paragraph


def _norm_key(s: str) -> str:
    """Create normalized key for snippet (remove extra spaces, limit length)."""
    return " ".join((s or "").split())[:400]


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("DOCX Translator App")
        self.setGeometry(100, 100, 1200, 800)

        # --- UI ---
        self.original_view = QWebEngineView()
        self.original_view.selectionChanged.connect(self.update_selection)
        self.translated_view = QWebEngineView()

        splitter = QSplitter()
        splitter.addWidget(self.original_view)
        splitter.addWidget(self.translated_view)
        splitter.setSizes([600, 600])

        central = QWidget()
        layout = QVBoxLayout()

        # --- Language selector ---
        lang_layout = QHBoxLayout()
        lang_label = QLabel("Target language:")
        self.lang_combo = QComboBox()
        langs = [
            ("English", "en"),
            ("Vietnamese", "vi"),
            ("French", "fr"),
            ("Spanish", "es"),
            ("German", "de"),
            ("Chinese (Simplified)", "zh-CN"),
        ]
        for name, code in langs:
            self.lang_combo.addItem(name, code)
        self.lang_combo.setCurrentIndex(1)
        lang_layout.addWidget(lang_label)
        lang_layout.addWidget(self.lang_combo)
        layout.addLayout(lang_layout)

        # --- Buttons ---
        btn_layout = QHBoxLayout()
        upload_btn = QPushButton("Upload DOCX")
        upload_btn.clicked.connect(self.upload_file)
        translate_btn = QPushButton("Translate selection")
        translate_btn.clicked.connect(self.translate_selected)
        download_btn = QPushButton("Download translated file")
        download_btn.clicked.connect(self.download_file)
        btn_layout.addWidget(upload_btn)
        btn_layout.addWidget(translate_btn)
        btn_layout.addWidget(download_btn)
        layout.addLayout(btn_layout)

        layout.addWidget(splitter)
        central.setLayout(layout)
        self.setCentralWidget(central)

        # --- State ---
        self.original_doc: Document | None = None
        self.translated_doc: Document | None = None
        self.selection_start = -1
        self.selection_end = -1

        self.segment_map: dict[str, dict] = {}

        # Scroll-sync state
        self._pending_scroll_y_right = None

        try:
            self.translator = TranslatorColumns("translate-tool.json")
        except Exception as e:
            QMessageBox.critical(self, "Translator error", str(e))
            self.translator = None

    # ---------- Upload DOCX ----------
    def upload_file(self):
        fname, _ = QFileDialog.getOpenFileName(self, "Select DOCX file", "", "DOCX (*.docx)")
        if not fname:
            return

        self.original_doc = Document(fname)

        buffer = BytesIO()
        self.original_doc.save(buffer)
        buffer.seek(0)
        self.translated_doc = Document(buffer)

        self.segment_map.clear()

        html1, base1 = self.translator.docx_to_html(self.original_doc)
        self.original_view.setHtml(html1, QUrl.fromLocalFile(base1 + "/"))

        html2, base2 = self.translator.docx_to_html(self.translated_doc)
        self.translated_view.setHtml(html2, QUrl.fromLocalFile(base2 + "/"))

    # ---------- Selection ----------
    def update_selection(self):
        js = """
        (function() {
            var sel = window.getSelection();
            if (sel.rangeCount > 0) {
                var r = sel.getRangeAt(0);
                var pre = r.cloneRange();
                pre.selectNodeContents(document.body);
                pre.setEnd(r.startContainer, r.startOffset);
                var start = pre.toString().length;
                var end = start + r.toString().length;
                return [start,end];
            }
            return [-1,-1];
        })();
        """
        self.original_view.page().runJavaScript(js, self.handle_selection_result)

    def handle_selection_result(self, result):
        if result and len(result) == 2:
            self.selection_start, self.selection_end = result
        else:
            self.selection_start = self.selection_end = -1

    # ---------- Helper: find selection span across multiple paragraphs ----------
    def _find_selection_span_across_paragraphs(self, selected_text: str):
        """
        Tìm vị trí selected_text (đã strip) trong toàn bộ tài liệu (nối paragraphs bằng ' ').
        Trả về (first_idx, start_off, last_idx, end_off) nếu tìm thấy; ngược lại None.
        """
        if not self.translated_doc:
            return None

        par_texts = [p.text for p in self.translated_doc.paragraphs]
        sep = " "
        full_text = sep.join(par_texts)

        sel = (selected_text or "").strip()
        if not sel:
            return None

        def _norm(s: str) -> str:
            return " ".join(s.split())

        norm_full = _norm(full_text)
        norm_sel = _norm(sel)

        start = norm_full.find(norm_sel)
        if start == -1:
            return None
        end = start + len(norm_sel)

        # Map offset lại theo đoạn gốc
        acc = 0
        first_idx = start_off = None
        for i, t in enumerate(par_texts):
            norm_t = _norm(t)
            acc_end = acc + len(norm_t)
            if start <= acc_end:
                first_idx = i
                start_off = max(start - acc, 0)
                break
            acc = acc_end + 1  # cộng thêm khoảng trắng phân tách

        acc = 0
        last_idx = end_off = None
        for j, t in enumerate(par_texts):
            norm_t = _norm(t)
            acc_end = acc + len(norm_t)
            if end <= acc_end:
                last_idx = j
                end_off = max(end - acc, 0)
                break
            acc = acc_end + 1

        if first_idx is None or last_idx is None:
            return None

        return (first_idx, start_off, last_idx, end_off)

    def _split_selection_by_paragraphs(self, span, selected_text: str):
        """
        Từ span (first_idx, start_off, last_idx, end_off) -> list các mảnh theo từng paragraph:
        trả về list[(para_idx, start_pos, end_pos, original_piece_text)].
        """
        first_idx, start_off, last_idx, end_off = span
        parts = []

        paragraphs = self.translated_doc.paragraphs

        if first_idx == last_idx:
            p = paragraphs[first_idx]
            piece = p.text[start_off:end_off]
            parts.append((first_idx, start_off, end_off, piece))
            return parts

        # paragraph đầu
        p0 = paragraphs[first_idx]
        piece0 = p0.text[start_off:]
        parts.append((first_idx, start_off, len(p0.text), piece0))

        # các paragraph giữa
        for i in range(first_idx + 1, last_idx):
            pi = paragraphs[i]
            piecei = pi.text
            if piecei:
                parts.append((i, 0, len(pi.text), piecei))

        # paragraph cuối
        pk = paragraphs[last_idx]
        piecek = pk.text[:end_off]
        parts.append((last_idx, 0, end_off, piecek))

        return parts

    # ---------- Scroll sync helpers ----------
    def _reload_translated_view_preserve_scroll(self):
        """Reload right pane and keep its scroll aligned with left pane."""
        if not self.translated_doc:
            return
        html, base = self.translator.docx_to_html(self.translated_doc)

        js_get_scroll_y = "Math.round(window.pageYOffset || document.documentElement.scrollTop || document.body.scrollTop || 0);"

        def _got_left_y(y):
            try:
                self.translated_view.loadFinished.disconnect(self._on_right_load_finished_set_scroll)
            except Exception:
                pass
            self._pending_scroll_y_right = int(y) if isinstance(y, (int, float)) else 0
            self.translated_view.loadFinished.connect(self._on_right_load_finished_set_scroll)
            self.translated_view.setHtml(html, QUrl.fromLocalFile(base + "/"))

        self.original_view.page().runJavaScript(js_get_scroll_y, _got_left_y)

    def _on_right_load_finished_set_scroll(self, ok):
        y = int(self._pending_scroll_y_right or 0)
        self.translated_view.page().runJavaScript(f"window.scrollTo(0, {y});")
        self._pending_scroll_y_right = None
        try:
            self.translated_view.loadFinished.disconnect(self._on_right_load_finished_set_scroll)
        except Exception:
            pass

    # ---------- Translate selected ----------
    def translate_selected(self):
        if not self.translator:
            return
        if self.selection_start == -1 or self.selection_end <= self.selection_start:
            QMessageBox.warning(self, "Warning", "Please highlight some text on the left pane.")
            return

        self.original_view.page().runJavaScript(
            "window.getSelection().toString();",
            self._translate_callback
        )

    def _translate_callback(self, selected_text: str):
        if not selected_text or not selected_text.strip():
            return

        key = _norm_key(selected_text)
        try:
            target_lang = self.lang_combo.currentData()

            replaced = False
            para_idx_used = None

            # --- Case 1: Thay trong 1 paragraph ---
            for idx, p in enumerate(self.translated_doc.paragraphs):
                pos = p.text.find(selected_text.strip())
                if pos != -1:
                    new_text = self.translator.translate_text(selected_text, target_lang)
                    replace_text_in_paragraph(p, pos, pos + len(selected_text.strip()), new_text)
                    replaced = True
                    para_idx_used = idx
                    self.segment_map[key] = {"para_idx": para_idx_used, "last_text": new_text}
                    break

            # --- Case 2: Dựa trên segment_map ---
            if not replaced and key in self.segment_map:
                info = self.segment_map[key]
                pi = info.get("para_idx")
                last = info.get("last_text", "")
                if isinstance(pi, int) and 0 <= pi < len(self.translated_doc.paragraphs) and last:
                    p = self.translated_doc.paragraphs[pi]
                    pos = p.text.find(last)
                    if pos != -1:
                        new_text = self.translator.translate_text(selected_text, target_lang)
                        replace_text_in_paragraph(p, pos, pos + len(last), new_text)
                        replaced = True
                        para_idx_used = pi
                        self.segment_map[key] = {"para_idx": para_idx_used, "last_text": new_text}

            # --- Case 3: Fallback multi-paragraph: dịch theo TỪNG ĐOẠN ---
            if not replaced:
                span = self._find_selection_span_across_paragraphs(selected_text)
                if span is not None:
                    parts = self._split_selection_by_paragraphs(span, selected_text)
                    if not parts:
                        QMessageBox.warning(self, "Not found",
                                            "Could not locate the segment to replace in the translated document.")
                        return

                    for (pi, s, e, piece_text) in parts:
                        if piece_text and piece_text.strip():
                            translated_piece = self.translator.translate_text(piece_text, target_lang)
                        else:
                            translated_piece = piece_text
                        replace_text_in_paragraph(self.translated_doc.paragraphs[pi], s, e, translated_piece)
                        if para_idx_used is None:
                            para_idx_used = pi

                    replaced = True
                    self.segment_map[key] = {"para_idx": para_idx_used, "last_text": ""}

            if not replaced:
                QMessageBox.warning(self, "Not found",
                                    "Could not locate the segment to replace in the translated document.")
                return

            # Reload right pane and keep scroll position aligned with the left
            self._reload_translated_view_preserve_scroll()

        except Exception as e:
            QMessageBox.critical(self, "Translation error", str(e))

    # ---------- Download DOCX ----------
    def download_file(self):
        if not self.translated_doc:
            QMessageBox.warning(self, "Warning", "No translated content to save.")
            return
        fname, _ = QFileDialog.getSaveFileName(self, "Save translated file", "", "DOCX (*.docx)")
        if fname:
            self.translated_doc.save(fname)
            QMessageBox.information(self, "Success", "Translated file has been saved.")

    def closeEvent(self, event):
        try:
            if self.translator and hasattr(self.translator, "cleanup_all_tmp"):
                self.translator.cleanup_all_tmp()
        finally:
            super().closeEvent(event)
